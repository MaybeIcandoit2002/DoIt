from docx import Document
from docx.shared import Cm
import random


#读取信息
fil = open('688.txt','rb')
l = fil.readlines()
lst = []
for i in range(len(l)):
    lst.append(l[i].decode(encoding = "utf-8").rstrip('\r\n'))
page = []
x = eval(input('请输入页码\n一次输入一页，结束输入0\n:'))
while x>0 and x<24:
    for i in range(eval(input('请输入需要多少张\n:'))):
        page.append(x)
    x = eval(input('请输入页码\n一次输入一页，结束输入0\n:'))


#创建文档
doc = Document()
doc0 = Document()
sec = doc.sections
sec1 = doc0.sections

def cel(p,q):
    sec0 = sec[p]
    sec0.top_margin = Cm(1)
    sec0.bottom_margin = Cm(2)
    sec0.left_margin = Cm(1.2)
    sec0.right_margin = Cm(1)
    sec2 = sec1[p]
    sec2.top_margin = Cm(1)
    sec2.bottom_margin = Cm(2)
    sec2.left_margin = Cm(1.2)
    sec2.right_margin = Cm(1)
    
    #添加数据
    num = 31
    t = []
    if q != 23 :
        doc.add_heading('第{}页,序号{}'.format(q,p+1))
        table = doc.add_table(num,2,style='Table Grid')
        doc0.add_heading('第{}页,序号{}'.format(q,p+1))
        table0 = doc0.add_table(num,2,style='Table Grid')
        E = lst[num*(q-1):num*q]
        C = lst[num*(q-1)+688:num*q+688]
        for i in range(num):
            t.append([E[i],C[i]])
        random.shuffle(t)
        for n in range(num):
            table.cell(n,0).text = t[n][0]
            table.cell(n,1).text = t[n][1]
            
            table0.cell(n,1).text = t[n][1]
    else:
        doc.add_heading('第{}页,序号{}'.format(q,p+1))
        table = doc.add_table(6,2,style='Table Grid')
        doc0.add_heading('第{}页,序号{}'.format(q,p+1))
        table0 = doc0.add_table(6,2,style='Table Grid')
        E = lst[num*(q-1):num*q-25]
        C = lst[num*(q-1)+688:num*q+688-25]
        for i in range(6):
            t.append([E[i],C[i]])
        random.shuffle(t)
        for n in range(6):
            table.cell(n,0).text = t[n][0]
            table.cell(n,1).text = t[n][1]

            table0.cell(n,1).text = t[n][1]
            
        
        
        
cel(0,page[0])
if len(page) != 1:
    for i in range(1,len(page)):
        doc.add_section()
        doc0.add_section()
        cel(i,page[i])

fil.close()
doc.save('答案.docx')
doc0.save('测试.docx')
